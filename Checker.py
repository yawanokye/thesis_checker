import re
import zipfile
from dataclasses import dataclass
from typing import List, Dict, Tuple, Optional

import pandas as pd
from docx import Document
from docx.shared import Pt
from lxml import etree

# =========================
# Models
# =========================
@dataclass
class Issue:
    rule_id: str
    severity: str   # High, Medium, Low
    message: str
    evidence: str
    location_hint: str            # e.g. "Paragraph 128"
    anchor_paragraph_index: int   # paragraph index to attach a Word comment


# =========================
# Guideline constants (from your document)
# =========================
# Front matter order (Dedication optional)
FRONT_MATTER_ORDER = [
    "Title Page or Inside Cover",   # we detect by heuristic, since text may differ
    "Declaration",
    "Abstract",
    "Acknowledgements",
    "Dedication",                   # optional
    "Table of Contents",
    "List of Tables",
    "List of Figures",
    "List of Acronyms",
]
# From guideline: preliminaries list includes title page, declaration, abstract, acknowledgements,
# dedication, table of contents, lists. :contentReference[oaicite:9]{index=9}
# Main body default is chapter one to five. :contentReference[oaicite:10]{index=10}
# Back matter references and appendices, plus PhD extras. :contentReference[oaicite:11]{index=11}

# Formatting requirements
REQ_FONT_SIZE = 12  # font size 12 :contentReference[oaicite:12]{index=12}
PREFERRED_FONT = "Times New Roman" # preferably Times New Roman :contentReference[oaicite:13]{index=13}
REQ_DOUBLE_SPACING = 2.0           # double spacing :contentReference[oaicite:14]{index=14}
REQ_LEFT_MARGIN_INCH = 2.0         # 5.08 cm = 2 inch :contentReference[oaicite:15]{index=15}
REQ_OTHER_MARGIN_INCH = 1.0        # 2.54 cm = 1 inch :contentReference[oaicite:16]{index=16}

# Abstract <= 1 page :contentReference[oaicite:17]{index=17}
# Keywords = 6 :contentReference[oaicite:18]{index=18}
# Dedication <= 2 lines :contentReference[oaicite:19]{index=19}

# Similarity Index Report submission mentioned in ethics/plagiarism section :contentReference[oaicite:20]{index=20}


# =========================
# APA Tables/Figures & APA citations
# =========================
APA_TABLE_LABEL_RE = re.compile(r"^\s*Table\s+(\d+)\s*$", re.IGNORECASE)
APA_FIGURE_LABEL_RE = re.compile(r"^\s*Figure\s+(\d+)\s*$", re.IGNORECASE)

BAD_DECIMAL_RE = re.compile(r"^\s*(Table|Figure)\s+\d+\.\d+", re.IGNORECASE)
BAD_CHAPTER_STYLE_RE = re.compile(r"^\s*(Table|Figure)\s+\d+(\.\d+)+", re.IGNORECASE)

# In-text citations (APA author-date practical subset)
PAREN_BLOCK_RE = re.compile(r"\(([^()]+?)\)")
APA_AUTHOR_YEAR_RE = re.compile(
    r"([A-Z][A-Za-z'’-]+(?:\s*&\s*[A-Z][A-Za-z'’-]+)?|[A-Z][A-Za-z'’-]+\s+et\s+al\.)\s*,\s*(\d{4})([a-b])?",
    re.IGNORECASE
)
NARRATIVE_RE = re.compile(r"\b([A-Z][A-Za-z'’-]+)\s*\(\s*(\d{4})([a-b])?\s*\)")


def iter_paragraphs(doc: Document) -> List[Tuple[int, str]]:
    return [(i, (p.text or "").strip()) for i, p in enumerate(doc.paragraphs)]


def find_exact_heading_index(paras: List[Tuple[int, str]], heading: str) -> Optional[int]:
    target = heading.strip().lower()
    for idx, text in paras:
        if text.lower() == target:
            return idx
    return None


def find_heading_like(paras: List[Tuple[int, str]], candidates: List[str]) -> Optional[int]:
    cand = {c.lower(): c for c in candidates}
    for idx, text in paras:
        t = text.lower()
        if t in cand:
            return idx
    return None


def extract_text_between(paras: List[Tuple[int, str]], start: int, end: Optional[int]) -> str:
    out = []
    for idx, t in paras:
        if idx < start:
            continue
        if end is not None and idx >= end:
            break
        out.append(t)
    return "\n".join(out).strip()


def normalise(s: str) -> str:
    s = (s or "").strip()
    s = s.replace("’", "'").replace("–", "-").replace("—", "-")
    s = re.sub(r"\s+", " ", s)
    return s.lower()


# =========================
# Rule group A: Structure and required sections
# =========================
def check_structure(doc: Document, degree_type: str) -> List[Issue]:
    issues: List[Issue] = []
    paras = iter_paragraphs(doc)

    # Core preliminaries in order
    # We detect these headings exactly as per guideline samples.
    prelim_required = ["Declaration", "Abstract", "Acknowledgements", "Table of Contents"]
    # Dedication optional but if present must be positioned after Acknowledgements
    # Lists: if tables/figures present, we enforce List of Tables/Figures later.

    indices: Dict[str, Optional[int]] = {}
    for h in prelim_required + ["Dedication", "List of Tables", "List of Figures", "List of Acronyms", "References", "Appendices"]:
        indices[h] = find_exact_heading_index(paras, h)

    # Check required preliminaries exist
    for h in prelim_required:
        if indices[h] is None:
            issues.append(Issue(
                rule_id="STR-FM-001",
                severity="High",
                message=f"Missing required front matter section: {h}.",
                evidence="Heading not found exactly.",
                location_hint="Front matter",
                anchor_paragraph_index=0
            ))

    # Order check among those that exist
    order = ["Declaration", "Abstract", "Acknowledgements", "Dedication", "Table of Contents"]
    existing = [(h, indices[h]) for h in order if indices[h] is not None]
    if existing:
        idxs = [i for _, i in existing]
        if idxs != sorted(idxs):
            issues.append(Issue(
                rule_id="STR-FM-002",
                severity="High",
                message="Front matter sections are not in the required order (Declaration → Abstract → Acknowledgements → Dedication (optional) → Table of Contents).",
                evidence=str(existing),
                location_hint="Front matter",
                anchor_paragraph_index=existing[0][1] if existing[0][1] is not None else 0
            ))

    # Main body chapters (default five-chapter structure)
    ch = [find_heading_like(paras, [f"Chapter {i}", f"CHAPTER {i}"]) for i in range(1, 6)]
    if any(x is None for x in ch):
        issues.append(Issue(
            rule_id="STR-CH-001",
            severity="High",
            message="Default structure expects Chapter One to Chapter Five headings.",
            evidence=f"Detected chapter heading indices: {ch}",
            location_hint="Main body",
            anchor_paragraph_index=(ch[0] or 0)
        ))
    else:
        if ch != sorted(ch):
            issues.append(Issue(
                rule_id="STR-CH-002",
                severity="High",
                message="Chapters are not ordered sequentially (Chapter 1–5).",
                evidence=f"Indices: {ch}",
                location_hint="Main body",
                anchor_paragraph_index=ch[0]
            ))

    # Back matter required
    if indices["References"] is None:
        issues.append(Issue(
            rule_id="STR-BM-001",
            severity="High",
            message="Missing References section.",
            evidence="Heading 'References' not found.",
            location_hint="Back matter",
            anchor_paragraph_index=(ch[-1] or 0)
        ))
    if indices["Appendices"] is None:
        issues.append(Issue(
            rule_id="STR-BM-002",
            severity="Medium",
            message="Appendices section not found. If you have appendices, add an 'Appendices' heading.",
            evidence="Heading 'Appendices' not found.",
            location_hint="Back matter",
            anchor_paragraph_index=(indices["References"] or (ch[-1] or 0))
        ))

    # PhD extras
    if degree_type.lower() == "phd":
        # VITA and copies of papers are specified for PhD only. :contentReference[oaicite:21]{index=21}
        vita = find_exact_heading_index(paras, "Vita")
        if vita is None:
            issues.append(Issue(
                rule_id="STR-PHD-001",
                severity="Medium",
                message="PhD theses should include a Vita section.",
                evidence="Heading 'Vita' not found.",
                location_hint="Back matter",
                anchor_paragraph_index=(indices["References"] or (ch[-1] or 0))
            ))

    return issues


# =========================
# Rule group: Abstract, Keywords, Dedication content rules
# =========================
def check_abstract_keywords_dedication(doc: Document) -> List[Issue]:
    issues: List[Issue] = []
    paras = iter_paragraphs(doc)

    abstract_idx = find_exact_heading_index(paras, "Abstract")
    kw_idx = find_heading_like(paras, ["Key Words", "Keywords"])
    ded_idx = find_exact_heading_index(paras, "Dedication")

    # Abstract <= 1 page: DOCX has no reliable "page" without Word layout engine.
    # MVP heuristic: abstract should be "short". We flag if it's clearly too long (e.g., > 600 words).
    if abstract_idx is not None:
        next_idx = min([i for i in [
            find_exact_heading_index(paras, "Key Words"),
            find_exact_heading_index(paras, "Acknowledgements"),
            find_exact_heading_index(paras, "Dedication"),
            find_exact_heading_index(paras, "Table of Contents"),
        ] if i is not None and i > abstract_idx] or [len(paras)])

        abstract_text = extract_text_between(paras, abstract_idx + 1, next_idx)
        wc = len([w for w in abstract_text.split() if w])
        if wc > 600:
            issues.append(Issue(
                rule_id="CNT-ABS-001",
                severity="High",
                message="Abstract appears longer than one page. Shorten it to not more than one page.",
                evidence=f"Estimated abstract word count: {wc}",
                location_hint=f"Paragraph {abstract_idx + 1}",
                anchor_paragraph_index=abstract_idx
            ))

    # Keywords: must be 6, listed vertically in alphabetical order. :contentReference[oaicite:22]{index=22}
    if kw_idx is not None:
        # capture next 12 lines/paragraphs until a blank or next heading-like
        kw_lines = []
        for i in range(kw_idx + 1, min(kw_idx + 30, len(paras))):
            t = paras[i][1]
            if not t:
                break
            if t.lower() in {"acknowledgements", "dedication", "table of contents", "chapter one", "chapter 1"}:
                break
            kw_lines.append(t)

        # allow comma-separated on one line, but guideline says vertical list, so we enforce vertical
        if len(kw_lines) == 1 and "," in kw_lines[0]:
            issues.append(Issue(
                rule_id="CNT-KW-001",
                severity="High",
                message="Key words should be listed vertically (one per line), not comma-separated.",
                evidence=kw_lines[0],
                location_hint=f"Paragraph {kw_idx + 1}",
                anchor_paragraph_index=kw_idx
            ))
            kws = [k.strip() for k in kw_lines[0].split(",") if k.strip()]
        else:
            kws = [t.strip("•- \t").strip() for t in kw_lines if t.strip()]

        if len(kws) != 6:
            issues.append(Issue(
                rule_id="CNT-KW-002",
                severity="High",
                message="Provide exactly six key words/phrases.",
                evidence=f"Detected {len(kws)}: {kws}",
                location_hint=f"Paragraph {kw_idx + 1}",
                anchor_paragraph_index=kw_idx
            ))

        # alphabetical check (case-insensitive)
        if kws and [k.lower() for k in kws] != sorted([k.lower() for k in kws]):
            issues.append(Issue(
                rule_id="CNT-KW-003",
                severity="Medium",
                message="Key words should be in alphabetical order.",
                evidence=str(kws),
                location_hint=f"Paragraph {kw_idx + 1}",
                anchor_paragraph_index=kw_idx
            ))

    # Dedication <= 2 lines :contentReference[oaicite:23]{index=23}
    if ded_idx is not None:
        # take next non-empty paragraph(s) until blank or next heading-like
        ded_lines = []
        for i in range(ded_idx + 1, min(ded_idx + 10, len(paras))):
            t = paras[i][1]
            if not t:
                break
            if t.lower() in {"table of contents", "chapter one", "chapter 1", "acknowledgements"}:
                break
            # treat each paragraph as one "line" for MVP
            ded_lines.append(t)

        if len(ded_lines) > 2:
            issues.append(Issue(
                rule_id="CNT-DED-001",
                severity="High",
                message="Dedication should not be longer than two lines.",
                evidence=f"Detected {len(ded_lines)} dedication lines/paragraphs.",
                location_hint=f"Paragraph {ded_idx + 1}",
                anchor_paragraph_index=ded_idx
            ))

    return issues


# =========================
# Rule group C: Formatting and layout
# =========================
def check_formatting(doc: Document) -> List[Issue]:
    issues: List[Issue] = []

    # Margins
    for sidx, section in enumerate(doc.sections):
        left = section.left_margin.inches
        right = section.right_margin.inches
        top = section.top_margin.inches
        bottom = section.bottom_margin.inches

        if abs(left - REQ_LEFT_MARGIN_INCH) > 0.05 or \
           abs(right - REQ_OTHER_MARGIN_INCH) > 0.05 or \
           abs(top - REQ_OTHER_MARGIN_INCH) > 0.05 or \
           abs(bottom - REQ_OTHER_MARGIN_INCH) > 0.05:
            issues.append(Issue(
                rule_id="FMT-MARG-001",
                severity="High",
                message="Margins must be 2-inch left and 1-inch top/bottom/right.",
                evidence=f"Section {sidx+1} margins (in): left={left:.2f}, right={right:.2f}, top={top:.2f}, bottom={bottom:.2f}",
                location_hint=f"Document sections: {sidx+1}",
                anchor_paragraph_index=0
            ))
            break  # one issue is enough

    # Font size and font name (sample across document)
    # We sample paragraphs and count violations
    font_bad = 0
    size_bad = 0
    sampled = 0

    for p in doc.paragraphs[:4000]:
        if not p.text.strip():
            continue
        for run in p.runs:
            if not run.text.strip():
                continue
            sampled += 1
            if run.font.size is not None:
                if abs(run.font.size.pt - REQ_FONT_SIZE) > 0.5:
                    size_bad += 1
            # if size is None, Word inherits; we don't penalise harshly

            if run.font.name is not None:
                if normalise(run.font.name) != normalise(PREFERRED_FONT):
                    font_bad += 1
            if sampled >= 500:
                break
        if sampled >= 500:
            break

    if size_bad > 40:
        issues.append(Issue(
            rule_id="FMT-FONT-001",
            severity="High",
            message="Font size should be 12 throughout (exceptions only for technical graphics).",
            evidence=f"Sample check: {size_bad} runs out of {sampled} differ from size 12.",
            location_hint="Main text sample",
            anchor_paragraph_index=0
        ))

    if font_bad > 80:
        issues.append(Issue(
            rule_id="FMT-FONT-002",
            severity="Medium",
            message="Font should preferably be Times New Roman throughout.",
            evidence=f"Sample check: {font_bad} runs out of {sampled} are not Times New Roman.",
            location_hint="Main text sample",
            anchor_paragraph_index=0
        ))

    # Double spacing and justification
    # Paragraph format: line_spacing and alignment
    spacing_bad = 0
    justify_bad = 0
    checked = 0

    for i, p in enumerate(doc.paragraphs[:3000]):
        t = p.text.strip()
        if not t:
            continue
        checked += 1

        # spacing: line_spacing can be None (inherit). We treat explicit non-2.0 as issue.
        ls = p.paragraph_format.line_spacing
        if ls is not None:
            # ls might be float or Length; python-docx returns float for "multiple"
            try:
                if isinstance(ls, float) and abs(ls - REQ_DOUBLE_SPACING) > 0.15:
                    spacing_bad += 1
            except Exception:
                pass

        # justify both sides. :contentReference[oaicite:24]{index=24}
        # Alignment enum: 3 is JUSTIFY in python-docx
        if p.paragraph_format.alignment is not None:
            if int(p.paragraph_format.alignment) != 3:
                justify_bad += 1

    if spacing_bad > 30:
        issues.append(Issue(
            rule_id="FMT-SPC-001",
            severity="High",
            message="Double spacing must be used throughout (long tables may be single-spaced).",
            evidence=f"Detected {spacing_bad} paragraphs with explicit non-double line spacing in a sample of {checked}.",
            location_hint="Main text sample",
            anchor_paragraph_index=0
        ))

    if justify_bad > 60:
        issues.append(Issue(
            rule_id="FMT-ALN-001",
            severity="Medium",
            message="Text should be justified on both left and right margins.",
            evidence=f"Detected {justify_bad} paragraphs with explicit non-justified alignment in a sample of {checked}.",
            location_hint="Main text sample",
            anchor_paragraph_index=0
        ))

    return issues


# =========================
# Rule group B: Chapter One template compliance (light MVP)
# =========================
def check_chapter_one_template(doc: Document) -> List[Issue]:
    issues: List[Issue] = []
    paras = iter_paragraphs(doc)

    ch1 = find_heading_like(paras, ["Chapter One", "CHAPTER ONE", "Chapter 1", "CHAPTER 1"])
    if ch1 is None:
        return issues

    # Heuristic: check presence of typical headings used in your preferred intro structure.
    required_subheads = [
        "Background of the Study",
        "Statement of the Problem",
        "Purpose of the Study",
        "Research Objectives",
        "Research Questions",
        "Significance of the Study",
        "Limitations of the Study",
        "Delimitation of the Study",
        "Structure of the Thesis",
    ]
    # search within next ~250 paragraphs or until Chapter Two
    ch2 = find_heading_like(paras, ["Chapter Two", "CHAPTER TWO", "Chapter 2", "CHAPTER 2"])
    end = ch2 if (ch2 is not None and ch2 > ch1) else min(len(paras), ch1 + 250)
    window = {t.lower(): idx for idx, t in paras[ch1:end] if t}

    missing = [h for h in required_subheads if h.lower() not in window]
    if missing:
        issues.append(Issue(
            rule_id="CH1-TPL-001",
            severity="High",
            message="Chapter One is missing expected sub-headings used in the approved structure.",
            evidence=f"Missing: {missing}",
            location_hint=f"Chapter One (around paragraph {ch1+1})",
            anchor_paragraph_index=ch1
        ))

    return issues


# =========================
# APA Tables/Figures: serial numbering + lists
# =========================
def detect_table_figure_labels(doc: Document) -> Dict[str, List[Dict]]:
    paras = iter_paragraphs(doc)
    tables, figures, bad = [], [], []

    for pidx, text in paras:
        if not text:
            continue
        if BAD_DECIMAL_RE.match(text) or BAD_CHAPTER_STYLE_RE.match(text):
            bad.append({"pidx": pidx, "text": text})
            continue
        mt = APA_TABLE_LABEL_RE.match(text)
        if mt:
            tables.append({"n": int(mt.group(1)), "pidx": pidx, "label": f"Table {mt.group(1)}"})
            continue
        mf = APA_FIGURE_LABEL_RE.match(text)
        if mf:
            figures.append({"n": int(mf.group(1)), "pidx": pidx, "label": f"Figure {mf.group(1)}"})
            continue
    return {"tables": tables, "figures": figures, "bad": bad}


def check_serial(items: List[Dict], kind: str, rule_id: str) -> List[Issue]:
    issues: List[Issue] = []
    if not items:
        return issues
    nums = [it["n"] for it in items]
    expected = list(range(1, max(nums) + 1))
    missing = [n for n in expected if n not in nums]
    dupes = sorted({n for n in nums if nums.count(n) > 1})

    if min(nums) != 1:
        issues.append(Issue(rule_id, "High", f"{kind} numbering must start at 1.", str(nums), "Captions", items[0]["pidx"]))
    if missing:
        issues.append(Issue(rule_id, "High", f"Missing {kind} numbers: {missing}. Use serial numbering only.", str(nums), "Captions", items[0]["pidx"]))
    if dupes:
        issues.append(Issue(rule_id, "High", f"Duplicate {kind} numbers: {dupes}.", str(nums), "Captions", items[0]["pidx"]))

    return issues


def check_lists(doc: Document, detected: Dict[str, List[Dict]]) -> List[Issue]:
    issues: List[Issue] = []
    paras = iter_paragraphs(doc)

    lot = find_exact_heading_index(paras, "List of Tables")
    lof = find_exact_heading_index(paras, "List of Figures")

    if detected["tables"] and lot is None:
        issues.append(Issue("APA-LOT-001", "High", "List of Tables is missing but tables are present.", f"{len(detected['tables'])} tables detected.", "Front matter", 0))
    if detected["figures"] and lof is None:
        issues.append(Issue("APA-LOF-001", "High", "List of Figures is missing but figures are present.", f"{len(detected['figures'])} figures detected.", "Front matter", 0))

    for b in detected["bad"]:
        issues.append(Issue("APA-CAP-001", "High", "Non-APA caption numbering detected. Use Table 1, Table 2… and Figure 1, Figure 2… only.", b["text"], f"Paragraph {b['pidx']+1}", b["pidx"]))

    return issues


# =========================
# APA citations + references match
# =========================
def parse_reference_key(entry: str) -> Optional[Tuple[str, str]]:
    m = re.match(r"^\s*([A-Z][A-Za-z'’-]+)\s*,", entry)
    if not m:
        return None
    surname = normalise(m.group(1))
    y = re.search(r"\((\d{4})([a-b])?\)", entry)
    if not y:
        return None
    year = (y.group(1) + (y.group(2) or "")).lower()
    return (surname, year)


def split_reference_entries(ref_text: str) -> List[str]:
    return [p.strip() for p in re.split(r"\n\s*\n", ref_text) if p.strip()]


def extract_intext_keys(full_text: str) -> List[Tuple[str, str]]:
    keys = []

    for m in NARRATIVE_RE.finditer(full_text):
        keys.append((normalise(m.group(1)), (m.group(2) + (m.group(3) or "")).lower()))

    for m in PAREN_BLOCK_RE.finditer(full_text):
        inside = m.group(1)
        chunks = [c.strip() for c in inside.split(";")]
        for c in chunks:
            mm = APA_AUTHOR_YEAR_RE.search(c)
            if not mm:
                continue
            lead = mm.group(1).split("&")[0].replace("et al.", "").strip()
            keys.append((normalise(lead), (mm.group(2) + (mm.group(3) or "")).lower()))

    return keys


def check_apa_only_citation_style(full_text: str) -> List[Issue]:
    issues: List[Issue] = []
    if re.search(r"\[\s*\d+\s*\]", full_text):
        issues.append(Issue("APA-CIT-001", "High", "Numeric citation style detected (e.g., [12]). Use APA author–date only.", "Found [number].", "Main text", 0))
    if re.search(r"\bibid\.|\bop\. cit\.", full_text, flags=re.IGNORECASE):
        issues.append(Issue("APA-CIT-002", "Medium", "Non-APA citation marker detected (ibid./op. cit.). Use APA author–date only.", "Found ibid./op. cit.", "Main text", 0))
    return issues


def check_references(doc: Document) -> List[Issue]:
    issues: List[Issue] = []
    paras = iter_paragraphs(doc)

    ref_idx = find_exact_heading_index(paras, "References")
    if ref_idx is None:
        return [Issue("APA-REF-001", "High", "References section not found.", "Heading 'References' not found.", "Back matter", 0)]

    # Build main text (exclude references block for fewer false positives)
    main_text = extract_text_between(paras, 0, ref_idx)

    # Extract references text after heading
    ref_text = extract_text_between(paras, ref_idx + 1, None)
    if not ref_text.strip():
        return [Issue("APA-REF-002", "High", "References section is present but empty.", "No entries after 'References'.", "References", ref_idx)]

    entries = split_reference_entries(ref_text)
    ref_keys: Dict[Tuple[str, str], int] = {}
    unparsed = 0
    for e in entries:
        k = parse_reference_key(e)
        if k is None:
            unparsed += 1
            continue
        ref_keys[k] = ref_keys.get(k, 0) + 1

    if unparsed > 0:
        issues.append(Issue("APA-REF-003", "Medium", "Some reference entries could not be parsed as APA author–date. Check APA consistency.", f"Unparsed: {unparsed}/{len(entries)}", "References", ref_idx))

    # In-text citation keys
    intext_keys = extract_intext_keys(main_text)
    issues.extend(check_apa_only_citation_style(main_text))

    if not intext_keys:
        issues.append(Issue("APA-CIT-010", "High", "No APA in-text citations detected in main text.", "No matches like (Surname, 2020) or Surname (2020).", "Main text", 0))

    missing = [k for k in intext_keys if k not in ref_keys]
    if missing:
        ex = ", ".join([f"{a} {y}" for a, y in missing[:10]])
        issues.append(Issue("APA-MATCH-001", "High", "Some in-text citations have no matching reference entry.", ex + (" ..." if len(missing) > 10 else ""), "Citations vs References", ref_idx))

    intext_set = set(intext_keys)
    uncited = [k for k in ref_keys.keys() if k not in intext_set]
    if uncited:
        ex = ", ".join([f"{a} {y}" for a, y in uncited[:10]])
        issues.append(Issue("APA-MATCH-002", "Medium", "Some reference entries do not appear to be cited in text.", ex + (" ..." if len(uncited) > 10 else ""), "Citations vs References", ref_idx))

    dups = [k for k, c in ref_keys.items() if c > 1]
    if dups:
        ex = ", ".join([f"{a} {y}" for a, y in dups[:10]])
        issues.append(Issue("APA-REF-010", "Medium", "Duplicate references detected (same lead author and year). Consider duplicates or missing a/b suffix.", ex, "References", ref_idx))

    return issues


# =========================
# Rule group E: Similarity Index Report mention
# =========================
def check_similarity_report_mention(doc: Document) -> List[Issue]:
    issues: List[Issue] = []
    text = "\n".join([p.text for p in doc.paragraphs if p.text]).lower()
    if "similarity index report" not in text and "turnitin" not in text:
        issues.append(Issue(
            rule_id="ETH-SIM-001",
            severity="Medium",
            message="No mention of Similarity Index Report found. Ensure it is submitted/attached where required.",
            evidence="Did not find 'Similarity Index Report' or 'Turnitin'.",
            location_hint="Document-wide",
            anchor_paragraph_index=0
        ))
    return issues


# =========================
# Word Comments injector (real comments in DOCX)
# =========================
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
}

def _qn(tag: str) -> str:
    prefix, local = tag.split(":")
    return f"{{{NS[prefix]}}}{local}"

def add_word_comments(input_docx: str, issues: List[Issue], output_docx: str) -> None:
    """
    Adds true Word comments (sidebar) to paragraphs indicated by issue.anchor_paragraph_index.
    Uses OOXML edits via zipfile + lxml.
    """
    # Group issues by paragraph index
    by_p: Dict[int, List[Issue]] = {}
    for iss in issues:
        by_p.setdefault(max(0, iss.anchor_paragraph_index), []).append(iss)

    with zipfile.ZipFile(input_docx, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    doc_xml = etree.fromstring(files["word/document.xml"])
    rels_xml = etree.fromstring(files["word/_rels/document.xml.rels"])

    # Ensure comments part exists
    comments_name = "word/comments.xml"
    if comments_name in files:
        comments_xml = etree.fromstring(files[comments_name])
        # get max comment id
        existing_ids = comments_xml.xpath("//w:comment/@w:id", namespaces=NS)
        next_id = max([int(x) for x in existing_ids], default=-1) + 1
    else:
        comments_xml = etree.Element(_qn("w:comments"), nsmap={"w": NS["w"]})
        next_id = 0

    # Ensure relationship to comments exists
    rels = rels_xml.xpath("//rel:Relationship", namespaces=NS)
    has_comments_rel = any(r.get("Type", "").endswith("/comments") for r in rels)
    if not has_comments_rel:
        # new rId
        existing_rids = []
        for r in rels:
            rid = r.get("Id", "")
            m = re.match(r"rId(\d+)", rid)
            if m:
                existing_rids.append(int(m.group(1)))
        new_rid = f"rId{(max(existing_rids) + 1) if existing_rids else 1}"
        rel_el = etree.SubElement(rels_xml, _qn("rel:Relationship"))
        rel_el.set("Id", new_rid)
        rel_el.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments")
        rel_el.set("Target", "comments.xml")

    # All paragraphs
    paras = doc_xml.xpath("//w:body/w:p", namespaces=NS)

    def make_comment_text(iss: Issue) -> str:
        return f"[{iss.severity}] {iss.rule_id}\n{iss.message}\nEvidence: {iss.evidence}\nWhere: {iss.location_hint}"

    # Insert comments
    for pidx, plist in by_p.items():
        if pidx >= len(paras):
            continue
        p = paras[pidx]

        # Find first run, else create one
        runs = p.xpath("./w:r", namespaces=NS)
        if runs:
            first_r = runs[0]
        else:
            first_r = etree.SubElement(p, _qn("w:r"))
            etree.SubElement(first_r, _qn("w:t")).text = ""

        # For each issue, add a comment anchored to first run
        for iss in plist:
            cid = str(next_id)
            next_id += 1

            # comments.xml entry
            c_el = etree.SubElement(comments_xml, _qn("w:comment"))
            c_el.set(_qn("w:id"), cid)
            c_el.set(_qn("w:author"), "Thesis Compliance Checker")
            c_el.set(_qn("w:initials"), "TCC")

            cp = etree.SubElement(c_el, _qn("w:p"))
            cr = etree.SubElement(cp, _qn("w:r"))
            ct = etree.SubElement(cr, _qn("w:t"))
            ct.text = make_comment_text(iss)

            # Anchor in document.xml: commentRangeStart, commentRangeEnd, commentReference
            start = etree.Element(_qn("w:commentRangeStart"))
            start.set(_qn("w:id"), cid)

            end = etree.Element(_qn("w:commentRangeEnd"))
            end.set(_qn("w:id"), cid)

            ref_run = etree.Element(_qn("w:r"))
            ref = etree.SubElement(ref_run, _qn("w:commentReference"))
            ref.set(_qn("w:id"), cid)

            # Insert start before first run, end after first run, then reference after end
            p.insert(p.index(first_r), start)
            p.insert(p.index(first_r) + 2, end)      # +2 because start inserted
            p.insert(p.index(end) + 1, ref_run)

    # Write back
    files["word/document.xml"] = etree.tostring(doc_xml, xml_declaration=True, encoding="UTF-8", standalone="yes")
    files["word/_rels/document.xml.rels"] = etree.tostring(rels_xml, xml_declaration=True, encoding="UTF-8", standalone="yes")
    files[comments_name] = etree.tostring(comments_xml, xml_declaration=True, encoding="UTF-8", standalone="yes")

    with zipfile.ZipFile(output_docx, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, content in files.items():
            zout.writestr(name, content)


# =========================
# Runner
# =========================
def run_checks(docx_path: str, degree_type: str) -> Tuple[List[Issue], Dict]:
    doc = Document(docx_path)
    issues: List[Issue] = []

    # A: Structure
    issues.extend(check_structure(doc, degree_type))

    # Content rules
    issues.extend(check_abstract_keywords_dedication(doc))

    # B: Chapter one template (MVP)
    issues.extend(check_chapter_one_template(doc))

    # C: Formatting
    issues.extend(check_formatting(doc))

    # Tables/Figures APA serial numbering and list presence
    detected = detect_table_figure_labels(doc)
    issues.extend(check_serial(detected["tables"], "Table", "APA-TBL-001"))
    issues.extend(check_serial(detected["figures"], "Figure", "APA-FIG-001"))
    issues.extend(check_lists(doc, detected))

    # APA references and matching
    issues.extend(check_references(doc))

    # Similarity Index mention
    issues.extend(check_similarity_report_mention(doc))

    # Score
    high = sum(1 for i in issues if i.severity == "High")
    med = sum(1 for i in issues if i.severity == "Medium")
    low = sum(1 for i in issues if i.severity == "Low")
    meta = {
        "issues_total": len(issues),
        "issues_high": high,
        "issues_medium": med,
        "issues_low": low
    }
    return issues, meta
